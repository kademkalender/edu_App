from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ── colour palette ──────────────────────────────────────────────
NAVY   = RGBColor(0x0A, 0x19, 0x29)
GOLD   = RGBColor(0xFF, 0xC1, 0x07)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT  = RGBColor(0xF5, 0xF7, 0xFA)
DARK   = RGBColor(0x1A, 0x1A, 0x2E)
GREY   = RGBColor(0x55, 0x55, 0x55)

doc = Document()

# ── page margins ────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# ── default body font ───────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.font.color.rgb = DARK

# ════════════════════════════════════════════════════════════════
# HELPER FUNCTIONS
# ════════════════════════════════════════════════════════════════

def set_cell_bg(cell, hex_color: str):
    """Fill a table cell with a solid background colour."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)


def add_heading(text, level=1, color=NAVY):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if level == 1 else 10)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = color
    run.font.size = Pt({1: 20, 2: 15, 3: 13}.get(level, 12))
    run.font.name = 'Calibri'
    return p


def add_body(text, italic=False, color=DARK, size=11, space_after=6):
    p   = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name      = 'Calibri'
    run.font.size      = Pt(size)
    run.font.color.rgb = color
    run.italic         = italic
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_bullet(text, bold_prefix=None):
    p   = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold            = True
        r.font.name       = 'Calibri'
        r.font.size       = Pt(11)
        r.font.color.rgb  = DARK
    r2 = p.add_run(text)
    r2.font.name      = 'Calibri'
    r2.font.size      = Pt(11)
    r2.font.color.rgb = DARK
    p.paragraph_format.space_after = Pt(3)
    return p


def add_divider():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pb  = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '6')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), '0A1929')
    pb.append(bot)
    pPr.append(pb)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(10)


def styled_table(headers, rows, header_bg='0A1929', alt_bg='F5F7FA'):
    col_n = len(headers)
    tbl   = doc.add_table(rows=1 + len(rows), cols=col_n)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # header row
    hdr_cells = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        set_cell_bg(hdr_cells[i], header_bg)
        for run in hdr_cells[i].paragraphs[0].runs:
            run.bold            = True
            run.font.color.rgb  = WHITE
            run.font.size       = Pt(10.5)
            run.font.name       = 'Calibri'
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # data rows
    for r_idx, row_data in enumerate(rows):
        row_cells = tbl.rows[r_idx + 1].cells
        bg = alt_bg if r_idx % 2 == 0 else 'FFFFFF'
        for c_idx, cell_text in enumerate(row_data):
            row_cells[c_idx].text = cell_text
            set_cell_bg(row_cells[c_idx], bg)
            for run in row_cells[c_idx].paragraphs[0].runs:
                run.font.size      = Pt(10.5)
                run.font.name      = 'Calibri'
                run.font.color.rgb = DARK
            row_cells[c_idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    doc.add_paragraph()   # breathing room after table
    return tbl


# ════════════════════════════════════════════════════════════════
# COVER PAGE
# ════════════════════════════════════════════════════════════════

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title_p.add_run('TechEdu')
tr.bold            = True
tr.font.size       = Pt(38)
tr.font.color.rgb  = NAVY
tr.font.name       = 'Calibri'

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub_p.add_run('Mobile Learning Application')
sr.font.size      = Pt(18)
sr.font.color.rgb = GREY
sr.font.name      = 'Calibri'

doc.add_paragraph()
add_divider()
doc.add_paragraph()

info_lines = [
    ('Course',   'Mobile Application Development'),
    ('Group',    'Group 1'),
    ('Members',  'Muhammet Kadem Kalender  ·  Efe Erman  ·  Sarper Özkaya'),
    ('Platform', 'iOS  (Portrait & Landscape)'),
    ('Date',     datetime.date.today().strftime('%B %d, %Y')),
]
for label, value in info_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lb = p.add_run(f'{label}:  ')
    lb.bold            = True
    lb.font.size       = Pt(12)
    lb.font.color.rgb  = NAVY
    lb.font.name       = 'Calibri'
    vr = p.add_run(value)
    vr.font.size       = Pt(12)
    vr.font.color.rgb  = DARK
    vr.font.name       = 'Calibri'
    p.paragraph_format.space_after = Pt(5)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ════════════════════════════════════════════════════════════════

add_heading('1. Executive Summary', 1)
add_divider()

add_body(
    'TechEdu is an iOS mobile application developed for middle school students as part of a '
    'mobile application development course project. The application aims to make two key '
    'computer science topics — Operating System Structures and 5G Technology — accessible '
    'and engaging through a structured learning experience.',
)
add_body(
    'The app combines interactive lesson content with a gamified quiz system. Students '
    'progress through bite-sized lesson sections, then complete a timed multiple-choice quiz '
    'to test their understanding. Motivation is sustained through a points system, daily '
    'study streaks, and six unlockable badges. The application runs on both portrait and '
    'landscape orientations and follows Apple Human Interface Guidelines.',
)
add_body(
    'Built with React Native (Expo SDK 54) and React Native Reanimated 4, TechEdu delivers '
    'smooth animations and a consistent dark-themed UI across all screens.'
)


# ════════════════════════════════════════════════════════════════
# 2. PROJECT OVERVIEW
# ════════════════════════════════════════════════════════════════

add_heading('2. Project Overview', 1)
add_divider()

add_heading('2.1 Objectives', 2)
add_bullet('Deliver computer science fundamentals to middle school students in an approachable format.')
add_bullet('Replace passive reading with structured, section-by-section progression.')
add_bullet('Reinforce learning through immediate quiz feedback and detailed answer explanations.')
add_bullet('Encourage daily study habits via streaks, points, and badge milestones.')
add_bullet('Provide a polished, iOS-native visual experience in both portrait and landscape.')
doc.add_paragraph()

add_heading('2.2 Target Audience', 2)
add_body(
    'The primary users are middle school students (ages 11–14) who are beginning to engage '
    'with digital technology concepts. The UI uses simple language, large touch targets, '
    'and immediate visual feedback to accommodate users with limited prior experience in '
    'formal computer science education.'
)

add_heading('2.3 Topic Selection Rationale', 2)
styled_table(
    ['Topic', 'Rationale'],
    [
        ['Operating System Structures',
         'A cornerstone of computer science literacy — understanding how hardware and software '
         'interact through kernels, modes, and system calls gives students a mental model for '
         'all computing.'],
        ['5G Technology',
         'A timely, curiosity-driven topic that connects classroom learning to visible '
         'real-world infrastructure: smart cities, autonomous vehicles, IoT, and remote surgery.'],
    ]
)


# ════════════════════════════════════════════════════════════════
# 3. COURSE CONTENT
# ════════════════════════════════════════════════════════════════

add_heading('3. Course Content', 1)
add_divider()

add_heading('3.1 Lesson 1 — Operating System Structures', 2)
add_body(
    'This lesson introduces the foundational concepts behind every modern operating system. '
    'It is divided into five progressive sections, each building on the previous one.'
)
styled_table(
    ['#', 'Section Title', 'Key Concepts'],
    [
        ['1', 'What is an Operating System?',
         'Definition · core responsibilities (process, memory, file, I/O management) · '
         'examples: Windows, Linux, macOS, Android, iOS'],
        ['2', 'The Kernel',
         'Kernel role · scheduling · memory management · hardware drivers · '
         'kernel space vs user space'],
        ['3', 'User Mode vs Kernel Mode',
         'Two operating modes · hardware enforcement · why separation is critical for '
         'system stability and security'],
        ['4', 'System Calls',
         'Definition · five categories (process control, file management, device '
         'management, information, communication) · library abstraction (libc)'],
        ['5', 'OS Architectures',
         'Monolithic kernel (Linux) · Microkernel · Hybrid kernel (Windows, macOS) · '
         'performance vs reliability trade-offs'],
    ]
)
add_body('Quiz: 10 questions  ·  Passing threshold: 80 % (8 / 10)  ·  30-second timer per question',
         italic=True, color=GREY)
doc.add_paragraph()

add_heading('3.2 Lesson 2 — 5G Technology', 2)
add_body(
    'This lesson explores the fifth generation of mobile communication technology — from '
    'its core capabilities to the industries it is reshaping.'
)
styled_table(
    ['#', 'Section Title', 'Key Concepts'],
    [
        ['1', 'What is 5G?',
         'Generation timeline (1G → 5G) · "G" meaning · why 5G is an infrastructure '
         'shift, not just a speed upgrade'],
        ['2', 'Key Features',
         'eMBB (Enhanced Mobile Broadband) · URLLC (Ultra-Reliable Low-Latency '
         'Communications) · mMTC (Massive Machine-Type Communications)'],
        ['3', 'Latency',
         'Definition · 4G latency (30–50 ms) vs 5G latency (≤ 1 ms) · '
         'critical implications for autonomous vehicles and remote surgery'],
        ['4', 'Frequency Bands',
         'Low-band (< 1 GHz, wide coverage) · Mid-band (1–6 GHz, balance) · '
         'mmWave (> 24 GHz, multi-Gbps, short range)'],
        ['5', 'Use Cases',
         'Smart cities · autonomous vehicles · e-health & telemedicine · '
         'Industry 4.0 robot coordination · IoT (up to 1 M devices / km²)'],
    ]
)
add_body('Quiz: 10 questions  ·  Passing threshold: 80 % (8 / 10)  ·  30-second timer per question',
         italic=True, color=GREY)


# ════════════════════════════════════════════════════════════════
# 4. APPLICATION FEATURES
# ════════════════════════════════════════════════════════════════

add_heading('4. Application Features', 1)
add_divider()

add_heading('4.1 Authentication System', 2)
add_body(
    'The authentication module handles user registration and login without requiring a '
    'backend server. All data is stored locally on the device using AsyncStorage.'
)
add_bullet('New users register with a unique username and password.')
add_bullet('Passwords are hashed with SHA-256 before storage — plain text is never saved.')
add_bullet('Session tokens persist across app restarts; users stay logged in until they explicitly sign out.')
add_bullet('Logging out clears the session and returns the user to the Welcome screen.')
doc.add_paragraph()

add_heading('4.2 Lesson Experience', 2)
add_body(
    'The lesson flow is designed to maintain focus and prevent cognitive overload by '
    'presenting one section at a time.'
)
add_bullet('Each lesson is split into sections; students tap "Next" to advance.')
add_bullet('Three content block types: paragraph, callout box (key concept / tip / warning), and bullet list.')
add_bullet('A learning objective summary is shown at the top of each section.')
add_bullet('Navigation back to previous sections is supported for review.')
add_bullet('The quiz start button activates only after all sections are read.')
doc.add_paragraph()

add_heading('4.3 Quiz System', 2)
add_body(
    'The quiz is the primary assessment mechanism. Its design balances challenge with '
    'immediate, educational feedback.'
)
add_bullet('30-second countdown per question with a visual timer bar.')
add_bullet('Timer bar changes colour: green (> 20 s)  →  yellow (10–20 s)  →  red (< 10 s).')
add_bullet('Time expiry is treated as a wrong answer; the correct option is revealed.')
add_bullet('On selection: correct option turns green, wrong selection turns red.')
add_bullet('A detailed explanation text appears after each answered question.')
add_bullet('A progress bar at the top shows current position (e.g. Question 4 / 10).')
add_bullet('Option labels A / B / C / D are displayed in circular badges for clarity.')
doc.add_paragraph()

add_heading('4.4 Result Screen', 2)
add_body(
    'After the final question the student is taken to an animated result screen that '
    'summarises their performance and rewards.'
)
add_bullet('Animated score circle displaying X / 10 and percentage.')
add_bullet('Points earned and current streak shown side by side.')
add_bullet('Badge unlock notification with icon and description if a new badge is earned.')
add_bullet('Pass path: single "Go to Home" button.')
add_bullet('Fail path: "Retry Lesson" (returns to lesson start) and "Go to Home" buttons.')
doc.add_paragraph()

add_heading('4.5 Points & Gamification', 2)
styled_table(
    ['Trigger', 'Points'],
    [
        ['Each correct answer',          '+10'],
        ['Passing the quiz (≥ 8 / 10)',  '+50'],
        ['Perfect score (10 / 10)',       '+100'],
        ['First study session of the day', '+20'],
    ]
)

add_heading('4.6 Badge System', 2)
styled_table(
    ['Badge', 'Icon', 'Unlock Condition'],
    [
        ['First Step',     '🎯', 'Complete any quiz for the first time'],
        ['Passing Grade',  '✅', 'Pass a quiz with at least 8 / 10'],
        ['Perfect',        '⭐', 'Achieve a perfect score of 10 / 10'],
        ['3-Day Streak',   '🔥', 'Study on 3 consecutive days'],
        ['Weekly Hero',    '🏅', 'Study on 7 consecutive days'],
        ['Lesson Master',  '🏆', 'Complete 5 different lessons'],
    ]
)

add_heading('4.7 Profile Screen', 2)
add_body('The profile screen gives students an at-a-glance view of their learning progress.')
add_bullet('Username and total accumulated points.')
add_bullet('Current streak (consecutive days studied).')
add_bullet('Gallery of earned badges with icon and title.')
doc.add_paragraph()

add_heading('4.8 Landscape Mode Support', 2)
add_body(
    'Every screen in the application is fully usable when the device is rotated to landscape.'
)
add_bullet('All screens are wrapped in ScrollView — content that does not fit is scrollable rather than clipped.')
add_bullet('Content areas use maxWidth constraints (700 pt for lessons, 500–600 pt for other screens) and alignSelf: "center" to remain centered on wider canvases.')
add_bullet('SafeAreaView automatically protects left / right edges on notched iPhones in landscape.')
add_bullet('Quiz progress bar and timer bar stretch to full screen width in any orientation.')


# ════════════════════════════════════════════════════════════════
# 5. SCREEN FLOW & NAVIGATION
# ════════════════════════════════════════════════════════════════

add_heading('5. Screen Flow & Navigation', 1)
add_divider()

add_body(
    'The app uses React Navigation Native Stack, which mirrors iOS\'s standard '
    'push / pop navigation model and enables the native swipe-back gesture on every screen.'
)

styled_table(
    ['Screen', 'Route Name', 'Description'],
    [
        ['Welcome',  'Welcome',  'Entry point — Log In or Sign Up buttons'],
        ['Login',    'Login',    'Username + password form; navigates to Home on success'],
        ['Register', 'Register', 'New account creation; navigates to Home on success'],
        ['Home',     'Home',     'Lesson card list + streak badge + profile / logout buttons'],
        ['Lesson',   'Lesson',   'Section-by-section lesson reader with progress indicator'],
        ['Quiz',     'Quiz',     'Timed multiple-choice quiz with progress and timer bars'],
        ['Result',   'Result',   'Animated score, points, streak, and badge notifications'],
        ['Profile',  'Profile',  'User stats and full badge collection'],
    ]
)

add_body('Navigation rules:', italic=True)
add_bullet('The back button is hidden on the Home and Result screens to prevent accidental navigation.')
add_bullet('After login or registration the Welcome screen is removed from the stack (replace), so back-pressing does not return to it.')
add_bullet('Retrying a lesson from the Result screen uses replace to keep the stack shallow.')


# ════════════════════════════════════════════════════════════════
# 6. TARGET PLATFORM & DESIGN DECISIONS
# ════════════════════════════════════════════════════════════════

add_heading('6. Target Platform & Design Decisions', 1)
add_divider()

add_heading('6.1 Platform Choice — iOS', 2)
add_body(
    'The team chose iOS as the primary target platform. The application was developed and '
    'tested on iPhone using Expo Go. All design decisions were made with the Apple Human '
    'Interface Guidelines (HIG) in mind. The app also runs on Android via Expo, but '
    'Material Design conventions were not applied.'
)
doc.add_paragraph()

add_heading('6.2 HIG Compliance', 2)
styled_table(
    ['HIG Principle', 'Implementation in TechEdu'],
    [
        ['Touch target size (min 44 × 44 pt)',
         'All buttons and quiz option cards enforce minHeight: 64, well above the minimum.'],
        ['Typography hierarchy',
         'Four distinct font-size levels (title 28 pt, heading 20 pt, body 16 pt, caption 13 pt) '
         'ensure visual hierarchy on every screen.'],
        ['Stack navigation',
         'React Navigation Native Stack provides the iOS-native left-edge swipe-back gesture.'],
        ['Safe Area compliance',
         'SafeAreaView from react-native-safe-area-context handles notch, Dynamic Island, '
         'and landscape side insets.'],
        ['Spring animations',
         'Correct answer: withSpring bounce. Wrong answer: shake via withTiming + withSpring. '
         'Screen entries: FadeIn / FadeInDown / FadeInRight with staggered delays.'],
        ['Landscape support',
         'Centered max-width layouts and ScrollView wrappers ensure usability in both orientations.'],
    ]
)

add_heading('6.3 Visual Design', 2)
add_body('The application uses a dark navy and gold colour palette throughout:')
styled_table(
    ['Token', 'Hex', 'Usage'],
    [
        ['background',   '#0A1929', 'Screen backgrounds'],
        ['surface',      '#13294B', 'Cards, modals, table cells'],
        ['surfaceLight', '#1E3A5F', 'Hover / pressed states'],
        ['primary',      '#FFC107', 'Accent colour, links, active indicators'],
        ['success',      '#4CAF50', 'Correct answer highlight'],
        ['error',        '#F44336', 'Wrong answer highlight, timer warning'],
        ['textPrimary',  '#FFFFFF', 'Headings and body text'],
        ['textSecondary','#B0BEC5', 'Subtitles and helper text'],
    ]
)


# ════════════════════════════════════════════════════════════════
# 7. TECHNICAL ARCHITECTURE
# ════════════════════════════════════════════════════════════════

add_heading('7. Technical Architecture', 1)
add_divider()

add_heading('7.1 Technology Stack', 2)
styled_table(
    ['Technology', 'Version', 'Role'],
    [
        ['React Native',                    '0.81',  'Core cross-platform UI framework'],
        ['Expo SDK',                        '54',    'Build toolchain, OTA updates, native module access'],
        ['React Navigation (Native Stack)', '—',     'Screen routing with iOS-native transitions'],
        ['React Native Reanimated',         '4',     'Worklet-based animations running on the UI thread'],
        ['React Native Safe Area Context',  '—',     'Notch / Dynamic Island / landscape inset management'],
        ['AsyncStorage',                    '—',     'Persistent key-value store for users and progress'],
        ['Expo Status Bar',                 '—',     'Light status bar icons on the dark background'],
    ]
)

add_heading('7.2 Project Structure', 2)
add_body('The codebase is organised by feature responsibility:')
add_bullet('components/  — Four reusable UI components (Button, Input, LessonCard, QuizOption)')
add_bullet('context/     — AuthContext provides login state to the entire component tree via React Context')
add_bullet('data/        — lessons.js is the single source of truth for all lesson text and quiz questions')
add_bullet('navigation/  — AppNavigator configures the Native Stack with shared header styles')
add_bullet('screens/     — Eight screens, each a self-contained functional component')
add_bullet('theme/       — colors.js and spacing.js centralise all design tokens')
add_bullet('utils/       — constants.js (game rules), hash.js (SHA-256), storage.js (AsyncStorage helpers)')
doc.add_paragraph()

add_heading('7.3 State Management', 2)
add_body(
    'The app relies on React\'s built-in state primitives rather than an external state library:'
)
add_bullet('React Context (AuthContext) — global session state accessible to all screens.')
add_bullet('useState / useEffect — local quiz state (current question, selected answer, timer).')
add_bullet('useFocusEffect — refreshes home screen progress data each time the screen gains focus.')
add_bullet('useSharedValue / useAnimatedStyle — Reanimated shared values for UI-thread animations.')
doc.add_paragraph()

add_heading('7.4 Security Considerations', 2)
add_bullet('User passwords are never stored as plain text — SHA-256 hash is computed before write.')
add_bullet('All data remains local; no network requests are made, eliminating server-side attack surface.')
add_bullet('AsyncStorage keys are namespaced per username to prevent cross-user data leakage.')


# ════════════════════════════════════════════════════════════════
# 8. SETUP & RUNNING
# ════════════════════════════════════════════════════════════════

add_heading('8. Setup & Running the Application', 1)
add_divider()

add_heading('8.1 Requirements', 2)
add_bullet('Node.js 18 or higher installed on the development machine.')
add_bullet('iPhone with the Expo Go app installed (free on the App Store).')
add_bullet('Development machine and iPhone on the same Wi-Fi network.')
doc.add_paragraph()

add_heading('8.2 Installation Steps', 2)
p = doc.add_paragraph()
p.add_run('Step 1 — Install dependencies').bold = True
add_body('npm install', color=GREY, italic=True)

p = doc.add_paragraph()
p.add_run('Step 2 — Start the development server').bold = True
add_body('npx expo start', color=GREY, italic=True)
add_body(
    'A QR code appears in the terminal. Open the Camera app on iPhone, scan the code, '
    'and tap "Open in Expo Go". First load takes 2–3 minutes.'
)

p = doc.add_paragraph()
p.add_run('Optional — Run in a web browser').bold = True
add_body('npx expo start --web', color=GREY, italic=True)
doc.add_paragraph()

add_heading('8.3 Troubleshooting', 2)
styled_table(
    ['Error', 'Solution'],
    [
        ['"Project is incompatible with this version of Expo Go"',
         'Update Expo Go from the App Store — SDK 54 is required.'],
        ['Metro bundler / Babel crash',
         'Delete node_modules/.cache and .expo, then run: npx expo start --clear'],
        ['QR code not connecting',
         'Confirm both devices are on the same Wi-Fi network. Disable VPN if active.'],
    ]
)


# ════════════════════════════════════════════════════════════════
# 9. LIMITATIONS & FUTURE WORK
# ════════════════════════════════════════════════════════════════

add_heading('9. Limitations & Future Work', 1)
add_divider()

add_heading('9.1 Current Limitations', 2)
add_bullet('Only two lessons are available in the current version.')
add_bullet('User data is stored locally; there is no cloud sync or multi-device support.')
add_bullet('No instructor dashboard — teachers cannot view student progress.')
add_bullet('Quiz questions are fixed; there is no randomisation or question bank rotation.')
doc.add_paragraph()

add_heading('9.2 Proposed Enhancements', 2)
add_bullet('Expand the lesson library to cover additional computer science topics (networking, databases, algorithms).')
add_bullet('Add a backend with user accounts, cloud sync, and an instructor analytics dashboard.')
add_bullet('Implement question randomisation and difficulty-adaptive quiz logic.')
add_bullet('Support Turkish language alongside English with a language toggle.')
add_bullet('Add push notifications to remind students to maintain their daily streak.')


# ════════════════════════════════════════════════════════════════
# 10. CONCLUSION
# ════════════════════════════════════════════════════════════════

add_heading('10. Conclusion', 1)
add_divider()
add_body(
    'TechEdu demonstrates how a focused, well-structured mobile application can make '
    'abstract computer science concepts accessible to younger learners. By combining '
    'progressive lesson sections with timed quizzes, immediate feedback, and a gamified '
    'reward system, the app creates an environment where students are motivated to return '
    'daily and build lasting knowledge.'
)
add_body(
    'The technical stack — React Native, Expo SDK 54, and Reanimated 4 — enabled rapid '
    'development while delivering a native-quality user experience aligned with Apple\'s '
    'design principles. Landscape mode support and SHA-256 password hashing reflect the '
    'team\'s attention to both usability and security best practices.'
)
add_body(
    'This project establishes a solid foundation that can be extended with additional '
    'lessons, a cloud backend, and adaptive assessment features in future iterations.'
)

# ── save ────────────────────────────────────────────────────────
output_path = 'TechEdu_Project_Report.docx'
doc.save(output_path)
print(f'Report saved: {output_path}')
